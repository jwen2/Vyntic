"""Export helpers for workflow runs (Phase 4).

Exports intentionally use the persisted run snapshot rather than re-running
LLM work. Tabular exports produce a compact model-ready workbook; assistant
exports produce a simple Word memo from approved stage outputs.
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

from app.database import DocumentRow, SessionLocal
from app.models.workflow import Workflow
from app.models.workflow_run import TabularCell, WorkflowRun
from app.services.workflow_shapes import display_text, normalize_shape


def build_tabular_xlsx(run: WorkflowRun, workflow: Workflow) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = _sheet_title(workflow.name)

    columns = sorted(workflow.columns, key=lambda c: c.order_index)
    row_keys = _run_row_keys(run)
    doc_names = _doc_name_map(run.deal_id)

    ws.append(["Document" if workflow.row_source == "one_doc_per_row" else "Question"] + [c.label for c in columns])
    header_fill = PatternFill("solid", fgColor="1F2937")
    for cell in ws[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill

    cells_by_key = {(cell.row_key, cell.column_id): cell for cell in run.cells}
    for row_key in row_keys:
        label = doc_names.get(row_key, row_key) if workflow.row_source == "one_doc_per_row" else row_key
        row = [label]
        for column in columns:
            row.append(_export_value(cells_by_key.get((row_key, column.id))))
        ws.append(row)

    ws.freeze_panes = "B2"
    for idx, _ in enumerate(columns, start=2):
        ws.column_dimensions[get_column_letter(idx)].width = 18
    ws.column_dimensions["A"].width = 36 if workflow.row_source == "one_doc_per_row" else 52

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_assistant_docx(run: WorkflowRun, workflow: Workflow) -> bytes:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading(workflow.name, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Run #{run.run_number}").bold = True
    meta.add_run(f" · {len(run.document_ids)} document{'s' if len(run.document_ids) != 1 else ''} analyzed")

    for stage in sorted(run.stage_outputs, key=lambda s: s.order_index):
        body = (stage.edited_md if stage.edited_md is not None else stage.output_md).strip()
        if not body:
            continue
        doc.add_heading(f"{stage.order_index}. {stage.label}", level=1)
        _add_markdownish_body(doc, body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_export_filename(name: str, run_number: int, ext: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_") or "workflow"
    return f"{slug}_run_{run_number}.{ext}"


def _run_row_keys(run: WorkflowRun) -> list[str]:
    keys: list[str] = []
    for cell in run.cells:
        if cell.row_key not in keys:
            keys.append(cell.row_key)
    return keys or run.document_ids


def _doc_name_map(deal_id: str) -> dict[str, str]:
    db = SessionLocal()
    try:
        rows = db.query(DocumentRow).filter(DocumentRow.deal_id == deal_id).all()
        return {row.doc_id: row.filename for row in rows}
    finally:
        db.close()


def _export_value(cell: TabularCell | None) -> Any:
    """Spreadsheet-facing value for a cell.

    Numeric metrics export as real numbers so Excel can sum them; every other
    shape goes through the shared `display_text` flattener rather than a
    private copy of it.
    """
    if cell is None or cell.status == "error":
        return ""
    shape = normalize_shape(cell.answer_formatted)
    if shape is None:
        return _strip_sources(cell.answer)
    if shape["kind"] == "metric" and not shape.get("raw") and shape.get("value") is not None:
        return shape["value"]
    return display_text(shape, compact=True)


def _strip_sources(value: str) -> str:
    return re.sub(r"\[Source\s+\d+\]", "", value or "", flags=re.IGNORECASE).strip()


def _add_markdownish_body(doc: Document, body: str) -> None:
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("#"):
            text = line.lstrip("#").strip()
            if text:
                doc.add_heading(text, level=2)
            continue
        if line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        doc.add_paragraph(line)


def _sheet_title(name: str) -> str:
    cleaned = re.sub(r"[:\\/?*\[\]]", " ", name).strip() or "Workflow Run"
    return cleaned[:31]
