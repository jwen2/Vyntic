"""
IC Report generator — builds a professional Word document from pre-computed
workstream results using python-docx.
"""
import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.models.report import ICReportRequest, ReportCitation

# ── Styling constants ──────────────────────────────────────────────────────

FONT_NAME = "Calibri"
TITLE_COLOR = RGBColor(0x1B, 0x2A, 0x4A)  # dark navy
HEADING_COLOR = RGBColor(0x1F, 0x39, 0x6C)
ACCENT_COLOR = RGBColor(0x2B, 0x57, 0x9A)
MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)
BODY_COLOR = RGBColor(0x1F, 0x2A, 0x37)

SOURCE_PATTERN = re.compile(r"\[Source\s+(\d+)\]")


def _apply_font(run, size: int, color: RGBColor = BODY_COLOR, bold: bool = False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _add_page_number(section):
    """Add page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page number field
    run = p.add_run()
    _apply_font(run, 8, MUTED_COLOR)
    fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = run2._element.makeelement(qn("w:instrText"), {})
    instrText.text = " PAGE "
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fldChar2)


def _add_footer_text(section, text: str):
    """Add text line to footer above page number."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _apply_font(run, 7, MUTED_COLOR)


def _add_title_page(doc: Document, title: str, deal_names: list[str]):
    """Create a clean title page."""
    # Spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, 0, 0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _apply_font(run, 28, TITLE_COLOR, bold=True)
    _set_paragraph_spacing(p, 0, 12)

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    _apply_font(run, 10, ACCENT_COLOR)
    _set_paragraph_spacing(p, 6, 18)

    # Deal names
    for name in deal_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        _apply_font(run, 14, HEADING_COLOR)
        _set_paragraph_spacing(p, 2, 2)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, 24, 0)
    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    run = p.add_run(date_str)
    _apply_font(run, 11, MUTED_COLOR)

    # Confidentiality notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, 36, 0)
    run = p.add_run("CONFIDENTIAL — For Investment Committee Use Only")
    _apply_font(run, 9, MUTED_COLOR, bold=True)

    doc.add_page_break()


def _add_executive_summary(doc: Document, request: ICReportRequest):
    """Add executive summary table with deal overview."""
    p = doc.add_paragraph()
    run = p.add_run("Executive Summary")
    _apply_font(run, 20, TITLE_COLOR, bold=True)
    _set_paragraph_spacing(p, 0, 12)

    # Summary description
    p = doc.add_paragraph()
    num_deals = len(request.deals)
    workstream_set: set[str] = set()
    for deal in request.deals:
        for ws in deal.workstreams:
            workstream_set.add(ws.workstream_name)
    ws_list = ", ".join(sorted(workstream_set)) if workstream_set else "N/A"
    run = p.add_run(
        f"This report covers {num_deals} deal{'s' if num_deals != 1 else ''} "
        f"across the following workstreams: {ws_list}."
    )
    _apply_font(run, 10, BODY_COLOR)
    _set_paragraph_spacing(p, 0, 12)

    # Deal overview table
    headers = ["Deal", "Stage", "Sector Tags", "Documents", "Workstreams Included"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        _apply_font(run, 9, RGBColor(0xFF, 0xFF, 0xFF), bold=True)

    # Data rows
    for deal in request.deals:
        row = table.add_row()
        ws_names = [ws.workstream_name for ws in deal.workstreams]
        values = [
            deal.name,
            deal.stage or "—",
            ", ".join(deal.tags) if deal.tags else "—",
            str(deal.document_count),
            ", ".join(ws_names) if ws_names else "—",
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _apply_font(run, 9, BODY_COLOR)

    doc.add_paragraph()  # spacing
    doc.add_page_break()


def _build_citation_index(citations: list[ReportCitation]) -> dict[int, ReportCitation]:
    """Build a 1-indexed map from source number to citation."""
    return {i + 1: c for i, c in enumerate(citations)}


def _is_table_row(line: str) -> bool:
    """Check if a line looks like a markdown table row."""
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 3


def _is_separator_row(line: str) -> bool:
    """Check if a table row is a separator (e.g. |---|---|)."""
    cells = [c.strip() for c in line.strip("|").split("|")]
    return all(re.match(r"^:?-+:?$", c) for c in cells if c)


def _parse_table_row(line: str) -> list[str]:
    """Split a markdown table row into cell values."""
    return [c.strip() for c in line.strip("|").split("|")]


def _is_delta_value(text: str) -> str | None:
    """Detect if a cell value is a positive/negative delta."""
    cleaned = re.sub(r"\*+", "", text).strip()
    cleaned = SOURCE_PATTERN.sub("", cleaned).strip()
    if not cleaned:
        return None
    if cleaned.startswith("+"):
        return "positive"
    if re.match(r"^[-\u2013]", cleaned) and not re.match(r"^[-\u2013:]+$", cleaned):
        return "negative"
    if "(" in cleaned and re.search(r"\([-\u2013]", cleaned):
        return "negative"
    return None


# Table cell colors for deltas
_POSITIVE_COLOR = RGBColor(0x05, 0x6B, 0x2E)  # dark green
_NEGATIVE_COLOR = RGBColor(0xB9, 0x1C, 0x1C)  # dark red
_POSITIVE_BG = "E8F5E9"
_NEGATIVE_BG = "FFEBEE"
_HEADER_BG = "334155"  # slate-700
_ALT_ROW_BG = "F8FAFC"  # slate-50


def _add_word_table(
    doc: Document,
    header_cells: list[str],
    data_rows: list[list[str]],
    citation_map: dict[int, ReportCitation],
):
    """Create a professionally styled Word table from parsed markdown table data."""
    if not header_cells or not data_rows:
        return

    num_cols = len(header_cells)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove default style and apply manual formatting for full control
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(header_cells):
        cell = hdr_row.cells[i]
        cell.text = ""
        # Background color
        shading = cell._element.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): _HEADER_BG,
        })
        cell._element.get_or_add_tcPr = lambda: cell._element.find(qn("w:tcPr")) or cell._element.makeelement(qn("w:tcPr"), {})
        tcPr = cell._element.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = cell._element.makeelement(qn("w:tcPr"), {})
            cell._element.insert(0, tcPr)
        tcPr.append(shading)

        p = cell.paragraphs[0]
        clean = re.sub(r"\*+", "", header).strip()
        run = p.add_run(clean)
        _apply_font(run, 9, RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        _set_paragraph_spacing(p, 2, 2, 1.0)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        is_alt = row_idx % 2 == 1

        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            cell.text = ""

            # Detect delta coloring
            delta = _is_delta_value(cell_text)
            bg_color = None
            text_color = BODY_COLOR

            if delta == "positive":
                bg_color = _POSITIVE_BG
                text_color = _POSITIVE_COLOR
            elif delta == "negative":
                bg_color = _NEGATIVE_BG
                text_color = _NEGATIVE_COLOR
            elif is_alt:
                bg_color = _ALT_ROW_BG

            if bg_color:
                tcPr = cell._element.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = cell._element.makeelement(qn("w:tcPr"), {})
                    cell._element.insert(0, tcPr)
                shading = cell._element.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): bg_color,
                })
                tcPr.append(shading)

            p = cell.paragraphs[0]
            _set_paragraph_spacing(p, 1, 1, 1.0)

            # Render cell content with bold handling and footnotes
            clean = cell_text.strip()
            _render_text_with_footnotes(p, clean, citation_map, font_size=9, color=text_color)

    doc.add_paragraph()  # spacing after table


def _render_text_with_footnotes(
    paragraph,
    text: str,
    citation_map: dict[int, ReportCitation],
    font_size: int = 10,
    color: RGBColor = BODY_COLOR,
):
    """Render text into a paragraph, handling **bold**, [Source N] footnotes."""
    parts = SOURCE_PATTERN.split(text)

    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            if part:
                bold_parts = re.split(r"\*\*(.*?)\*\*", part)
                for bi, bp in enumerate(bold_parts):
                    if not bp:
                        continue
                    if bi % 2 == 1:
                        run = paragraph.add_run(bp)
                        _apply_font(run, font_size, color, bold=True)
                    else:
                        run = paragraph.add_run(bp)
                        _apply_font(run, font_size, color)
        else:
            source_num = int(part)
            citation = citation_map.get(source_num)
            if citation:
                snippet = citation.text_snippet[:200]
                if len(citation.text_snippet) > 200:
                    snippet += "..."
                footnote_text = (
                    f'{citation.source_file}, p. {citation.page} — '
                    f'"{snippet}"'
                )
                _add_footnote(paragraph, footnote_text)
            else:
                run = paragraph.add_run(f" [{source_num}]")
                _apply_font(run, 7, ACCENT_COLOR)


def _add_answer_with_footnotes(
    doc: Document,
    answer: str,
    citations: list[ReportCitation],
):
    """Render answer text, replacing [Source N] with Word footnotes.

    Markdown tables are parsed into styled Word tables with delta coloring.
    Each footnote contains: source_file, p. page_number — "snippet..."
    """
    citation_map = _build_citation_index(citations)

    lines = answer.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Skip markdown heading markers
        if line.startswith("#"):
            line = line.lstrip("#").strip()

        # ── Markdown table detection ──
        if _is_table_row(line):
            # Collect all contiguous table rows
            table_lines: list[str] = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse: first row = header, skip separator, rest = data
            if len(table_lines) >= 2:
                header_cells = _parse_table_row(table_lines[0])
                data_rows: list[list[str]] = []
                for tl in table_lines[1:]:
                    if _is_separator_row(tl):
                        continue
                    data_rows.append(_parse_table_row(tl))

                _add_word_table(doc, header_cells, data_rows, citation_map)
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, 2, 4, 1.15)

        # Bullet point
        if line.startswith("- ") or line.startswith("* "):
            line = line[2:]
            p.style = doc.styles["List Bullet"]

        _render_text_with_footnotes(p, line, citation_map)
        i += 1


def _add_footnote(paragraph, text: str):
    """Insert a Word footnote at the current position in the paragraph.

    Uses low-level oxml because python-docx doesn't expose footnotes natively.
    """
    # Get or create footnotes part
    doc_part = paragraph.part
    footnotes_part = _get_or_create_footnotes_part(doc_part)

    # Determine next footnote ID
    footnotes_element = footnotes_part._element
    existing = footnotes_element.findall(qn("w:footnote"))
    # IDs 0 and 1 are reserved for separator and continuation
    next_id = max((int(fn.get(qn("w:id"), 0)) for fn in existing), default=1) + 1

    # Build footnote XML
    from lxml import etree

    footnote = etree.SubElement(
        footnotes_element, qn("w:footnote"), {qn("w:id"): str(next_id)}
    )
    fn_para = etree.SubElement(footnote, qn("w:p"))

    # Paragraph properties for footnote text style
    fn_pPr = etree.SubElement(fn_para, qn("w:pPr"))
    fn_pStyle = etree.SubElement(fn_pPr, qn("w:pStyle"))
    fn_pStyle.set(qn("w:val"), "FootnoteText")

    # Footnote reference mark inside the footnote itself
    fn_run_ref = etree.SubElement(fn_para, qn("w:r"))
    fn_rPr_ref = etree.SubElement(fn_run_ref, qn("w:rPr"))
    fn_rStyle = etree.SubElement(fn_rPr_ref, qn("w:rStyle"))
    fn_rStyle.set(qn("w:val"), "FootnoteReference")
    etree.SubElement(fn_run_ref, qn("w:footnoteRef"))

    # Space after ref mark
    fn_run_space = etree.SubElement(fn_para, qn("w:r"))
    fn_t_space = etree.SubElement(fn_run_space, qn("w:t"))
    fn_t_space.set(qn("xml:space"), "preserve")
    fn_t_space.text = " "

    # Footnote body text
    fn_run_text = etree.SubElement(fn_para, qn("w:r"))
    fn_rPr_text = etree.SubElement(fn_run_text, qn("w:rPr"))
    fn_sz = etree.SubElement(fn_rPr_text, qn("w:sz"))
    fn_sz.set(qn("w:val"), "16")  # 8pt
    fn_szCs = etree.SubElement(fn_rPr_text, qn("w:szCs"))
    fn_szCs.set(qn("w:val"), "16")
    fn_color = etree.SubElement(fn_rPr_text, qn("w:color"))
    fn_color.set(qn("w:val"), "6B7280")
    fn_t = etree.SubElement(fn_run_text, qn("w:t"))
    fn_t.set(qn("xml:space"), "preserve")
    fn_t.text = text

    # Now insert a footnote reference in the main paragraph
    run = paragraph.add_run()
    run_element = run._element
    rPr = run_element.makeelement(qn("w:rPr"), {})
    rStyle = rPr.makeelement(qn("w:rStyle"), {qn("w:val"): "FootnoteReference"})
    rPr.append(rStyle)
    run_element.insert(0, rPr)

    footnote_ref = run_element.makeelement(
        qn("w:footnoteReference"), {qn("w:id"): str(next_id)}
    )
    run_element.append(footnote_ref)


class _XmlPart:
    """Minimal Part-like wrapper whose blob always reflects the live lxml tree.

    python-docx's save walks parts via `part.rels`, so we expose an empty
    Relationships collection and the handful of attributes the serializer needs.
    """

    def __init__(self, partname, content_type, element, package):
        from docx.opc.rel import Relationships
        self.partname = partname
        self.content_type = content_type
        self._element = element
        self.package = package
        self.rels = Relationships(partname.baseURI)

    def before_marshal(self):
        pass

    @property
    def blob(self):
        from lxml import etree
        return etree.tostring(self._element, xml_declaration=True, encoding="UTF-8", standalone=True)


def _get_or_create_footnotes_part(doc_part):
    """Get the footnotes part from the document, creating if it doesn't exist."""
    from lxml import etree

    # Check if footnotes part already exists
    for rel in doc_part.rels.values():
        if "footnotes" in rel.reltype:
            tp = rel.target_part
            # Ensure _element is parsed for existing parts
            if not hasattr(tp, "_element") or tp._element is None:
                tp._element = etree.fromstring(tp.blob)
            return tp

    # Create footnotes part from scratch
    from docx.opc.packuri import PackURI

    footnotes_xml = (
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:footnote w:type="separator" w:id="0">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="1">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )

    element = etree.fromstring(footnotes_xml.encode("utf-8"))

    footnotes_part = _XmlPart(
        PackURI("/word/footnotes.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        element,
        doc_part.package,
    )

    doc_part.relate_to(
        footnotes_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    )

    return footnotes_part


def generate_ic_report(request: ICReportRequest) -> io.BytesIO:
    """Generate a Word document IC report and return it as a BytesIO buffer."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    font.color.rgb = BODY_COLOR

    # ── Heading styles ──
    for level, (size, color) in {
        1: (18, TITLE_COLOR),
        2: (14, HEADING_COLOR),
        3: (11, ACCENT_COLOR),
    }.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_NAME
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True

    # ── Footer with page numbers ──
    _add_footer_text(section, "Confidential — Generated by Vyntic")
    _add_page_number(section)

    # ── Title page ──
    deal_names = [d.name for d in request.deals]
    _add_title_page(doc, request.title, deal_names)

    # ── Executive Summary ──
    _add_executive_summary(doc, request)

    # ── Deal sections ──
    for deal_idx, deal in enumerate(request.deals):
        # Deal heading
        doc.add_heading(deal.name, level=1)

        # Deal metadata line
        meta_parts = []
        if deal.stage:
            meta_parts.append(f"Stage: {deal.stage}")
        if deal.tags:
            meta_parts.append(f"Sectors: {', '.join(deal.tags)}")
        if deal.document_count:
            meta_parts.append(f"Documents: {deal.document_count}")
        if meta_parts:
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(meta_parts))
            _apply_font(run, 9, MUTED_COLOR)
            _set_paragraph_spacing(p, 0, 8)

        # Workstream sections
        for ws in deal.workstreams:
            doc.add_heading(f"{ws.workstream_name}", level=2)

            for q in ws.questions:
                # Question as heading 3
                doc.add_heading(q.label, level=3)

                # Question text (italic, smaller)
                p = doc.add_paragraph()
                run = p.add_run(q.question)
                _apply_font(run, 9, MUTED_COLOR)
                run.italic = True
                _set_paragraph_spacing(p, 0, 4)

                # Answer with footnoted citations
                _add_answer_with_footnotes(doc, q.answer, q.citations)

        # Page break between deals (not after last)
        if deal_idx < len(request.deals) - 1:
            doc.add_page_break()

    # ── Citation Appendix ──
    doc.add_page_break()
    doc.add_heading("Source Documents Referenced", level=1)

    # Collect all unique sources across the report
    seen_sources: dict[tuple[str, int], ReportCitation] = {}
    for deal in request.deals:
        for ws in deal.workstreams:
            for q in ws.questions:
                for c in q.citations:
                    key = (c.source_file, c.page)
                    if key not in seen_sources:
                        seen_sources[key] = c

    if seen_sources:
        # Source table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Source File", "Page", "Excerpt"]):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            _apply_font(run, 9, RGBColor(0xFF, 0xFF, 0xFF), bold=True)

        for (source_file, page), citation in sorted(seen_sources.items()):
            row = table.add_row()
            snippet = citation.text_snippet[:150]
            if len(citation.text_snippet) > 150:
                snippet += "..."
            for i, val in enumerate([source_file, str(page), snippet]):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                _apply_font(run, 8, BODY_COLOR)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No source citations were included in this report.")
        _apply_font(run, 10, MUTED_COLOR)

    # ── Save to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
